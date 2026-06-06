import json
import re
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any
from zipfile import ZipFile
import xml.etree.ElementTree as ET

import pandas as pd
import pdfplumber
from docx import Document
from openpyxl import load_workbook


ROOT_DIR = Path(r"D:\2026\zhiyuan\高考志Y系列资料")
INVENTORY_JSON = Path(r"D:\2026\zhiyuan\data_work\resource_inventory.json")
OUTPUT_JSON = Path(r"D:\2026\zhiyuan\data_work\admission_candidate_files.json")
FEASIBILITY_MD = Path(r"D:\2026\zhiyuan\docs\admission_data_feasibility_report.md")
PLAN_MD = Path(r"D:\2026\zhiyuan\docs\score_search_data_plan.md")

CORE_TYPES = {"录取统计", "一分一段表", "招生计划", "专业选科要求"}


def sanitize(value: Any) -> str:
    if value is None:
        return ""
    text = str(value).replace("\u3000", " ").strip()
    return re.sub(r"\s+", " ", text)


def excel_probe(path: Path) -> dict[str, Any]:
    probe: dict[str, Any] = {"sheets": [], "header_keywords": [], "sample_rows": []}
    try:
        excel = pd.ExcelFile(path)
        probe["sheets"] = excel.sheet_names
        header_counter: Counter[str] = Counter()
        sample_rows: list[dict[str, Any]] = []
        for sheet_name in excel.sheet_names[:5]:
            df = pd.read_excel(path, sheet_name=sheet_name, nrows=5, header=None)
            rows = [[sanitize(cell) for cell in row] for row in df.fillna("").values.tolist()]
            if rows:
                sample_rows.append({"sheet": sheet_name, "rows": rows[:5]})
                for cell in rows[0]:
                    if cell:
                        header_counter[cell] += 1
        probe["header_keywords"] = [item for item, _ in header_counter.most_common(20)]
        probe["sample_rows"] = sample_rows[:3]
    except Exception as exc:
        probe["error"] = f"{type(exc).__name__}: {exc}"
    return probe


def pdf_probe(path: Path) -> dict[str, Any]:
    probe: dict[str, Any] = {"page_count": 0, "sample_pages": [], "needs_ocr": False}
    try:
        with pdfplumber.open(path) as pdf:
            probe["page_count"] = len(pdf.pages)
            for index, page in enumerate(pdf.pages[:2], start=1):
                text = sanitize(page.extract_text() or "")
                probe["sample_pages"].append({"page": index, "text": text[:1200]})
            probe["needs_ocr"] = all(not page["text"] for page in probe["sample_pages"])
    except Exception as exc:
        probe["error"] = f"{type(exc).__name__}: {exc}"
    return probe


def docx_probe(path: Path) -> dict[str, Any]:
    probe: dict[str, Any] = {"paragraphs": [], "tables": []}
    try:
        doc = Document(path)
        paragraphs = [sanitize(p.text) for p in doc.paragraphs if sanitize(p.text)]
        probe["paragraphs"] = paragraphs[:8]
        for table in doc.tables[:3]:
            rows = []
            for row in table.rows[:5]:
                rows.append([sanitize(cell.text) for cell in row.cells])
            probe["tables"].append(rows)
    except Exception as exc:
        probe["error"] = f"{type(exc).__name__}: {exc}"
    return probe


def pptx_probe(path: Path) -> dict[str, Any]:
    probe: dict[str, Any] = {"slides": []}
    ns = {
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "p": "http://schemas.openxmlformats.org/presentationml/2006/main",
    }
    try:
        with ZipFile(path) as archive:
            slide_names = sorted(
                name for name in archive.namelist() if name.startswith("ppt/slides/slide") and name.endswith(".xml")
            )
            for slide_name in slide_names[:3]:
                with archive.open(slide_name) as handle:
                    root = ET.parse(handle).getroot()
                texts = [sanitize(node.text) for node in root.findall(".//a:t", ns) if sanitize(node.text)]
                probe["slides"].append({"slide": slide_name, "texts": texts[:20]})
    except Exception as exc:
        probe["error"] = f"{type(exc).__name__}: {exc}"
    return probe


def probe_file(path: Path) -> dict[str, Any]:
    ext = path.suffix.lower()
    if ext in {".xls", ".xlsx"}:
        return excel_probe(path)
    if ext == ".pdf":
        return pdf_probe(path)
    if ext == ".docx":
        return docx_probe(path)
    if ext == ".pptx":
        return pptx_probe(path)
    return {"note": "unsupported for deep probe"}


def classify_probe(record: dict[str, Any], probe: dict[str, Any]) -> tuple[str, str]:
    doc_type = record["likely_document_type"]
    ext = record["extension"]
    if "error" in probe:
        return "低", f"探测报错：{probe['error']}"
    if doc_type in {"录取统计", "一分一段表", "招生计划"} and ext in {".xls", ".xlsx"}:
        headers = " ".join(probe.get("header_keywords", []))
        if any(keyword in headers for keyword in ["分数", "位次", "院校", "专业", "计划", "最低"]):
            return "高", "Excel 表头已出现分数/位次/院校/专业/计划等关键字段。"
        return "中", "Excel 可读，但关键字段需人工进一步确认。"
    if doc_type == "专业选科要求":
        if ext in {".xls", ".xlsx"}:
            return "高", "如果是表格版选科要求，结构化成本较低。"
        if ext == ".pdf":
            return "中", "PDF 版选科要求可抽文本，但可能仍需清洗。"
    if ext == ".pdf":
        if probe.get("needs_ocr"):
            return "低", "前两页基本无文本，疑似扫描版或图片版 PDF。"
        text = " ".join(page.get("text", "") for page in probe.get("sample_pages", []))
        if any(keyword in text for keyword in ["院校", "专业组", "最低分", "位次", "分数段"]):
            return "中", "PDF 抽样页包含核心字段，可尝试分省分批抽取。"
        return "低", "PDF 抽样页未稳定出现可结构化字段。"
    if ext == ".docx":
        combined = " ".join(probe.get("paragraphs", []))
        if "分段" in combined or "位次" in combined or probe.get("tables"):
            return "中", "DOCX 有文本或表格，可做小规模抽取。"
        return "低", "DOCX 更像说明材料。"
    if ext == ".pptx":
        return "低", "PPTX 不适合作为主数据源。"
    return "未知", "暂未形成稳定结论。"


def load_inventory() -> list[dict[str, Any]]:
    return json.loads(INVENTORY_JSON.read_text(encoding="utf-8"))


def write_reports(candidates: list[dict[str, Any]]) -> None:
    by_type: dict[str, list[dict[str, Any]]] = {doc_type: [] for doc_type in CORE_TYPES}
    for item in candidates:
        by_type[item["likely_document_type"]].append(item)

    def top_items(doc_type: str, limit: int = 20) -> list[dict[str, Any]]:
        items = sorted(
            by_type[doc_type],
            key=lambda x: (
                x.get("structured_feasibility") != "高",
                x["likely_year"],
                x["likely_province"],
                x["file_name"],
            ),
        )
        return items[:limit]

    lines = [
        "# 录取数据可行性评估",
        "",
        f"- 评估时间：`{datetime.now().isoformat(timespec='seconds')}`",
        f"- 候选文件总数：`{len(candidates)}`",
        "",
        "## 哪些文件最适合生成 admissions.json",
        "",
    ]
    for item in top_items("录取统计"):
        lines.append(
            f"- `{item['file_path']}` | `{item['likely_year']}` | `{item['likely_province']}` | `{item['likely_subject_type']}` | `{item['structured_feasibility']}` | {item['feasibility_reason']}"
        )
    lines.extend(["", "## 哪些文件适合生成 rank_tables.json", ""])
    for item in top_items("一分一段表"):
        lines.append(
            f"- `{item['file_path']}` | `{item['likely_year']}` | `{item['likely_province']}` | `{item['likely_subject_type']}` | `{item['structured_feasibility']}` | {item['feasibility_reason']}"
        )
    lines.extend(["", "## 哪些文件适合生成 plans.json", ""])
    for item in top_items("招生计划"):
        lines.append(
            f"- `{item['file_path']}` | `{item['likely_year']}` | `{item['likely_province']}` | `{item['structured_feasibility']}` | {item['feasibility_reason']}"
        )
    lines.extend(["", "## 哪些文件适合补 subject_requirement", ""])
    for item in top_items("专业选科要求"):
        lines.append(
            f"- `{item['file_path']}` | `{item['likely_year']}` | `{item['structured_feasibility']}` | {item['feasibility_reason']}"
        )

    pdf_ocr = [item for item in candidates if item["extension"] == ".pdf" and item["probe"].get("needs_ocr")]
    lines.extend(["", "## 哪些 PDF 需要 OCR", ""])
    if pdf_ocr:
        for item in pdf_ocr[:40]:
            lines.append(f"- `{item['file_path']}`")
    else:
        lines.append("- 当前抽样候选文件中没有明显扫描版核心 PDF。")

    complete = [item for item in candidates if item["structured_feasibility"] == "高"]
    weak = [item for item in candidates if item["structured_feasibility"] == "低"]
    lines.extend(
        [
            "",
            "## 哪些文件字段最完整",
            "",
            *(f"- `{item['file_path']}`" for item in complete[:20]),
            "",
            "## 哪些文件字段缺失严重",
            "",
            *(f"- `{item['file_path']}` | {item['feasibility_reason']}" for item in weak[:20]),
            "",
            "## 是否可以先做单省单科类试点",
            "",
            "- 可以。优先建议 `山西 2025 物理类/历史类` 或 `山东 2025 夏季高考文化成绩` 试点。",
            "- 原因：一分一段和投档/最低分文件命名清晰，按省按科类拆分明显，后续抽取边界最容易控制。",
            "",
        ]
    )
    FEASIBILITY_MD.write_text("\n".join(lines), encoding="utf-8")

    plan_lines = [
        "# 按分数查询志愿数据方案",
        "",
        "## 1. 目前是否具备“按分数查询志愿”的最低数据条件？",
        "",
        "- 部分具备，但还不具备全国统一上线条件。",
        "- 如果按“单省 + 单科类 + 单年份”试点，现有资料已经足够进入数据抽取验证阶段。",
        "",
        "## 2. 如果不具备，还缺哪些字段？",
        "",
        "- 同一省份下稳定统一的院校代码、院校专业组代码、专业代码映射。",
        "- 与录取数据同口径的一分一段表或位次表。",
        "- 对应年份招生计划数。",
        "- 选科要求与专业组之间的结构化映射。",
        "",
        "## 3. 最小可用版本应该先支持哪个省、哪个科类？",
        "",
        "- 建议先做 `山西 2025 物理类` 或 `山西 2025 历史类`。",
        "- 备选是 `山东 2025 夏季高考文化成绩`，但山东志愿规则和数据口径更特殊，适合作为第二试点。",
        "",
        "## 4. admissions.json 建议字段结构",
        "",
        "```json",
        "{",
        '  "province": "",',
        '  "year": 2025,',
        '  "subject_type": "物理类",',
        '  "batch": "本科",',
        '  "school_name": "",',
        '  "school_code": "",',
        '  "major_group_name": "",',
        '  "major_group_code": "",',
        '  "major_name": "",',
        '  "major_code": "",',
        '  "min_score": 0,',
        '  "min_rank": 0,',
        '  "avg_score": null,',
        '  "max_score": null,',
        '  "plan_count": null,',
        '  "source_file": "",',
        '  "source_page": null,',
        '  "source_row": null',
        "}",
        "```",
        "",
        "## 5. rank_tables.json 建议字段结构",
        "",
        "```json",
        "{",
        '  "province": "",',
        '  "year": 2025,',
        '  "subject_type": "物理类",',
        '  "score": 600,',
        '  "same_score_count": 0,',
        '  "cumulative_count": 0,',
        '  "rank_min": 0,',
        '  "rank_max": 0,',
        '  "source_file": "",',
        '  "source_row": null',
        "}",
        "```",
        "",
        "## 6. plans.json 建议字段结构",
        "",
        "```json",
        "{",
        '  "province": "",',
        '  "year": 2026,',
        '  "subject_type": "物理类",',
        '  "batch": "",',
        '  "school_name": "",',
        '  "school_code": "",',
        '  "major_group_name": "",',
        '  "major_group_code": "",',
        '  "major_name": "",',
        '  "major_code": "",',
        '  "plan_count": 0,',
        '  "tuition": null,',
        '  "duration": "",',
        '  "source_file": "",',
        '  "source_row": null',
        "}",
        "```",
        "",
        "## 7. 前端“按分数查志愿”页面需要哪些输入项",
        "",
        "- 省份",
        "- 年份",
        "- 科类（历史类 / 物理类 / 文科 / 理科 / 综合）",
        "- 分数",
        "- 可选：位次",
        "- 可选：批次",
        "- 可选：地区 / 院校标签 / 办学层次",
        "",
        "## 8. 哪些结果可以展示，哪些不能展示",
        "",
        "- 可以展示：历史最低分、历史最低位次、对应院校/专业组、招生计划、数据来源。",
        "- 可以展示：按历史数据筛出的“可参考院校列表”。",
        "- 不能展示：录取概率、稳录结论、保底承诺、拟录取判断。",
        "",
        "## 9. 如何避免把参考结果包装成录取承诺",
        "",
        "- 页面文案统一使用“历史数据参考”“辅助筛选”“仅供参考”。",
        "- 不使用“能上”“稳上”“保录”等确定性措辞。",
        "- 强制展示免责声明：请以当年省考试院、高校招生章程和招生计划为准。",
        "",
    ]
    PLAN_MD.write_text("\n".join(plan_lines), encoding="utf-8")


def main() -> None:
    inventory = load_inventory()
    candidates = [record for record in inventory if record["likely_document_type"] in CORE_TYPES]
    results = []
    for record in candidates:
        path = Path(record["file_path"])
        probe = probe_file(path)
        feasibility, reason = classify_probe(record, probe)
        results.append(
            {
                **record,
                "probe": probe,
                "structured_feasibility": feasibility,
                "feasibility_reason": reason,
            }
        )
    OUTPUT_JSON.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_JSON.write_text(json.dumps(results, ensure_ascii=False, indent=2), encoding="utf-8")
    write_reports(results)
    print(f"Probed {len(results)} candidate files into {OUTPUT_JSON}")


if __name__ == "__main__":
    main()
